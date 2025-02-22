import os
import pdfplumber
import docx
import fitz  # PyMuPDF
from docx import Document
import pypandoc 
from werkzeug.utils import secure_filename

ALLOWED_EXTENSIONS = {"pdf", "docx"}

def allowed_file(filename):
    """Check if the uploaded file has an allowed extension."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + "\n"
    return text.strip()

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    doc = docx.Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

def extract_text_from_pdf_layout(pdf_path):
    """Extract text from a PDF file layout"""
    text_data = []
    doc = fitz.open(pdf_path)
    for page in doc:
        blocks = page.get_text("blocks")  # Extracts text in block format
        for block in blocks:
            text_data.append(block[4])  # Extract text from block
    return "\n".join(text_data).strip() if text_data else None

def extract_text_from_docx_layout(docx_path):
    """Extract text from a docx file layout"""
    text_data = []
    doc = docx.Document(docx_path)
    for para in doc.paragraphs:
        text_data.append(para.text)
    return "\n".join(text_data).strip() if text_data else None

def process_resume_file(file, upload_folder,layout="false"):
    """Handles file saving and text extraction."""
    if file.filename == "" or not allowed_file(file.filename):
        return None, None

    filename = secure_filename(file.filename)
    file_path = os.path.join(upload_folder, filename)
    file.save(file_path)

    file_extension = filename.rsplit(".", 1)[1].lower()
    
    if file_extension == "pdf":
        if layout=="true":
            extracted_text = extract_text_from_pdf_layout(file_path)
        else:
            print("IN else",layout)
            extracted_text = extract_text_from_pdf(file_path)
        links = extract_links_from_pdf(file_path)
        print("PDF file detected")
        print("Extracted text:", extracted_text)
        print("Extracted links:", links)
    elif file_extension == "docx":
        print("DOCX file detected")
        if layout=="true":
            extracted_text = extract_text_from_docx_layout(file_path)
        else:
            extracted_text = extract_text_from_docx(file_path)
        links = extract_links_from_docx(file_path)
    else:
        return None, None

    return extracted_text, links, filename

def extract_links_from_pdf(pdf_path):
    links = []
    doc = fitz.open(pdf_path)

    for page in doc:
        for link in page.get_links():
            if "uri" in link:  # Extract only actual links
                links.append(link["uri"])

    return links

def extract_links_from_docx(docx_path):
    doc = Document(docx_path)
    links = []
    
    # Extract hyperlinks from the document
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype:
            links.append(rel.target_ref)
    
    return links

def create_doc(resume_text, filename="ATS_Optimized_Resume.docx"):
    print(type(resume_text))
    filename = filename.split(".")[0]+ ".pdf"
    # doc = Document()
    # doc.add_paragraph(resume_text)
    filepath = os.path.join("uploads", f"ATS_Optimized_{filename}")
    # with open(f"ATS_Optimized_{filename}.md", "w", encoding="utf-8") as file:
    #     file.write(resume_text)
        
    pypandoc.convert_text(resume_text, "pdf", outputfile=filepath, format="md",extra_args=[
            "--pdf-engine=pdflatex",
            "--variable", "geometry:a4paper,margin=1in",  # Adjust margins
            "--variable", "parskip=0.5em",  # Adjust paragraph spacing
            "--variable", "parindent=0pt"   # Remove paragraph indentation
        ])
    
    # doc.save(filepath)
    return filepath
